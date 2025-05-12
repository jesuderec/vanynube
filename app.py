# Archivo: app.py (Versi√≥n Completa con Correcci√≥n de Spinner)

import streamlit as st
import fitz  # PyMuPDF
import docx # Necesario para generar DOCX
import os
import google.generativeai as genai
import requests
from openai import OpenAI, APIError, AuthenticationError, RateLimitError, APIConnectionError # Importar OpenAI
import io # Necesario para crear DOCX en memoria

# --- 1. CONFIGURACI√ìN - Leer Claves SIEMPRE desde Variables de Entorno ---
# En Render, estas deben estar configuradas en el Dashboard -> Environment
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# --- PROMPT FIJO (Definido directamente aqu√≠) ---
PROMPT_FIJO = """
Rol: Eres un asistente de an√°lisis acad√©mico y apoyo docente.

Objetivo: Analizar el documento proporcionado para identificar posibles casos de plagio por copia textual, evaluar la aplicaci√≥n b√°sica de las normas de citaci√≥n APA (7¬™ edici√≥n), resumir los hallazgos, proponer una retroalimentaci√≥n espec√≠fica para el estudiante, y estimar la proporci√≥n de contenido original vs. citado/potencialmente plagiado para el docente.

Instrucciones:
analiza el siguiente documento de un estudiante universitario:
Realiza las siguientes tareas de forma compacta y utilizando terminos comunes para humanizar lomas posible la respuesta:

Retroalimentaci√≥n para el Estudiante
	Redactar en tono profesional, acad√©mico, constructivo y humano. Incluir:
o		Inicio: Se√±alar desviaciones en el uso de citas, referencias y normas APA vigentes.
o		Plagio/Integridad: Indicar si hay fragmentos tomados de fuentes externas sin aplicaci√≥n correcta de formato (uso de comillas o bloque APA), incluso si hay intento de citaci√≥n.
o		Uso de Inteligencia Regenerativa: Explicar si se detectan patrones de redacci√≥n t√≠picos de IA, indicando que el contenido generado √≠ntegramente por IA no es aceptable.
o		Importancia de la Originalidad: Reforzar la necesidad de producci√≥n propia y rigor acad√©mico.

Analsis de plagio
	Detecta Coincidencias Textuales: Identifica los fragmentos m√°s claros dentro del documento que parezcan ser copias textuales o casi textuales de fuentes externas.
	Analiza Cada Coincidencia Significativa: Para cada fragmento detectado:
	Texto del Estudiante: Cita el fragmento exacto.
	Fuente Probable: Indica la fuente original m√°s probable (formato APA).
	Intento de Citar: Menciona si existe cita en texto Y entrada en bibliograf√≠a.
	Uso de Comillas/Bloque (APA): Verifica si el texto copiado est√° correctamente formateado. Responde S√ç o NO.
	Evaluaci√≥n de Integridad: Si es NO, indica expl√≠citamente que la falta de comillas/bloque constituye una citaci√≥n incorrecta seg√∫n APA y es un indicador de posible plagio, incluso si se intent√≥ citar.

Revisa la Bibliograf√≠a:
	Formato General APA vigente: Eval√∫a de forma concisa si la lista de referencias sigue en general las normas APA vigente
	Consistencia: Se√±ala si faltan referencias para citas o viceversa.


Genera un Resumen del An√°lisis: 
	Sintetiza en un p√°rrafo los hallazgos principales del an√°lisis realizados

Estimaci√≥n de Contenido (Solo para el Docente): 
	Bas√°ndote en el an√°lisis anterior, proporciona una estimaci√≥n porcentual aproximada del contenido del documento:
		% Contenido Original / Elaboraci√≥n Propia: (Ej. % estimado de redacci√≥n, an√°lisis, conclusiones, justificaciones propias del estudiante).
		% Contenido Citado / Potencialmente Plagiado: (Ej. % estimado de copias textuales sin comillas, par√°frasis muy cercanas, informaci√≥n directamente extra√≠da de fuentes identificadas).

Consideraciones Finales:

	Prioriza la detecci√≥n de copias textuales sin formato adecuado.
	El resumen y la retroalimentaci√≥n deben basarse directamente en los hallazgos del an√°lisis.
	La estimaci√≥n porcentual  es exclusivamente informativa para el docente y no debe incluirse en la retroalimentaci√≥n al estudiante.
	El contexto es acad√©mico universitario (normas APA vigente.).
	Se debe generar el documento en parrafos concisos evitandoel uso de tablas """
# Podr√≠as definir un diccionario PROMPTS aqu√≠ si necesitas m√°s de uno
# PROMPTS = {"An√°lisis Acad√©mico": PROMPT_FIJO}
# DEFAULT_PROMPT_NAME = "An√°lisis Acad√©mico"

# --- Configuraci√≥n de p√°gina Streamlit ---
st.set_page_config(page_title="Analizador de IA - Gemini / DeepSeek / OpenAI", layout="wide")
st.title("üìÑ Analizador de IA (Gemini / DeepSeek / OpenAI)")

# --- Configurar Cliente Gemini (con verificaci√≥n) ---
gemini_model = None
if GOOGLE_API_KEY:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        gemini_model = genai.GenerativeModel('models/gemini-1.5-pro') # O 'gemini-1.5-flash'
        print("Cliente Gemini configurado.")
    except Exception as e:
        st.warning(f"Error al configurar Gemini (opci√≥n deshabilitada): {e}")
else:
    print("Advertencia: GOOGLE_API_KEY no encontrada. Opci√≥n Gemini deshabilitada.")
    # st.sidebar.warning("Gemini deshabilitado (falta API Key).") # Mover advertencias a la UI si se prefiere

# --- Verificar Claves OpenAI y DeepSeek (para feedback temprano) ---
if not OPENAI_API_KEY:
    print("Advertencia: OPENAI_API_KEY no encontrada. Opci√≥n OpenAI deshabilitada.")
    # st.sidebar.warning("OpenAI deshabilitado (falta API Key).")
else:
    print("Clave OpenAI encontrada.")

if not DEEPSEEK_API_KEY:
     print("Advertencia: DEEPSEEK_API_KEY no encontrada. Opci√≥n DeepSeek deshabilitada.")
     # st.sidebar.warning("DeepSeek deshabilitado (falta API Key).")
else:
     print("Clave DeepSeek encontrada.")


# --- Funciones de Llamada a APIs ---
# (Se mantienen aqu√≠ por simplicidad en este ejemplo, idealmente ir√≠an en utils/ai_operations.py)

def analizar_con_gemini(prompt):
    if not gemini_model: return "‚ö†Ô∏è Error: Cliente Gemini no inicializado."
    try:
        response = gemini_model.generate_content(prompt)
        if response.parts: return response.text
        else:
             reason = response.prompt_feedback.block_reason if response.prompt_feedback else 'Raz√≥n desconocida'
             return f"‚ö†Ô∏è Respuesta de Gemini bloqueada o vac√≠a. Raz√≥n: {reason}"
    except Exception as e: return f"‚ö†Ô∏è Error en llamada a Gemini: {e}"

def analizar_con_deepseek(prompt, temperature=0.7, top_p=0.9):
    if not DEEPSEEK_API_KEY: return "‚ö†Ô∏è Error: Clave API de DeepSeek no configurada."
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    data = {"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt}], "temperature": temperature, "top_p": top_p, "max_tokens": 4096}
    try:
        response = requests.post(url, headers=headers, json=data, timeout=180)
        response.raise_for_status()
        respuesta_json = response.json()
        if "choices" in respuesta_json and len(respuesta_json["choices"]) > 0 and "message" in respuesta_json["choices"][0] and "content" in respuesta_json["choices"][0]["message"]:
             return respuesta_json["choices"][0]["message"]["content"]
        else: return f"‚ö†Ô∏è Respuesta inesperada de DeepSeek: {respuesta_json}"
    except requests.exceptions.Timeout: return f"‚ö†Ô∏è Error de Timeout con DeepSeek (180s)."
    except requests.exceptions.RequestException as e: return f"‚ö†Ô∏è Error de conexi√≥n con DeepSeek: {e}."
    except (KeyError, IndexError, json.JSONDecodeError) as e: return f"‚ö†Ô∏è Error al procesar respuesta de DeepSeek: {e}"
    except Exception as e: return f"‚ö†Ô∏è Error inesperado con DeepSeek: {e}"

def analizar_con_openai(prompt, model_id="gpt-4o-mini"):
    if not OPENAI_API_KEY: return "‚ö†Ô∏è Error: Clave API de OpenAI no configurada."
    try:
        client = OpenAI(api_key=OPENAI_API_KEY, timeout=180.0)
        print(f"Llamando a OpenAI con el modelo: {model_id}")
        response = client.chat.completions.create(model=model_id, messages=[{"role": "user", "content": prompt}])
        if response.choices: return response.choices[0].message.content
        else: return "‚ö†Ô∏è OpenAI no devolvi√≥ 'choices' en la respuesta."
    except AuthenticationError: return f"‚ö†Ô∏è Error de Autenticaci√≥n OpenAI ({model_id}). Verifica API Key."
    except RateLimitError: return f"‚ö†Ô∏è Error de L√≠mite de Tasa OpenAI ({model_id}). Excedida la cuota."
    except APIConnectionError: return f"‚ö†Ô∏è Error de Conexi√≥n OpenAI ({model_id}). No se pudo conectar a la API."
    except APIError as e: return f"‚ö†Ô∏è Error en la API de OpenAI ({model_id}): {e}"
    except Exception as e: return f"‚ö†Ô∏è Error inesperado al llamar a OpenAI ({model_id}): {e}"

# --- Funciones de Archivo ---
def leer_archivo(file):
    try:
        extension = os.path.splitext(file.name)[1].lower()
        if extension == ".pdf":
            texto = ""
            pdf_bytes = file.getvalue()
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                for page in doc: texto += page.get_text()
            return texto
        elif extension == ".docx":
            doc = docx.Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif extension == ".txt":
            return file.getvalue().decode("utf-8", errors='ignore')
        else:
            st.warning(f"Tipo de archivo '{extension}' no soportado directamente.")
            return None
    except Exception as e:
        st.error(f"Error al leer el archivo '{file.name}': {e}")
        return None

def limpiar_texto(texto):
    if texto: return "\n".join(line.strip() for line in texto.splitlines() if line.strip())
    return None

# --- Funci√≥n para Nombre de Archivo DOCX ---
def generar_nombre_archivo_docx(nombre_archivo_original, modelo_usado):
  nombre_base, _ = os.path.splitext(os.path.basename(nombre_archivo_original))
  modelo_limpio = "".join(c for c in modelo_usado if c.isalnum() or c in ('-', '_')).replace('-', '').replace('_', '')
  if not modelo_limpio: modelo_limpio = "modelo"
  nuevo_nombre = f"{nombre_base}_{modelo_limpio}.docx"
  return nuevo_nombre

# --- Placeholder para Crear DOCX (¬°REEMPLAZAR CON TU L√ìGICA REAL!) ---
def crear_docx_bytes(texto_formateado):
    print("Generando DOCX (usando placeholder)...")
    document = docx.Document()
    document.add_paragraph(texto_formateado)
    # --- ¬°Aqu√≠ ir√≠a tu l√≥gica de formateo Markdown -> DOCX! ---
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- Estilos CSS ---
st.markdown(
    """
    <style>
    /* ... (Tu CSS aqu√≠ si lo necesitas) ... */
    .response-container { /* Clase para respuestas */
        font-size: 14px !important; border: 1px solid #eee; border-radius: 5px;
        padding: 15px; margin-bottom: 10px; background-color: #f9f9f9;
        white-space: pre-wrap; word-wrap: break-word; /* Para respetar saltos y ajustar palabras */
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- Inicializar st.session_state ---
# (Definir opciones disponibles basado en claves encontradas)
opciones_disponibles_runtime = []
if gemini_model: opciones_disponibles_runtime.append("Gemini")
if DEEPSEEK_API_KEY: opciones_disponibles_runtime.append("DeepSeek")
if OPENAI_API_KEY: opciones_disponibles_runtime.append("gpt-4o-mini") # A√±adir modelos OpenAI aqu√≠

default_modelos_runtime = [opciones_disponibles_runtime[0]] if opciones_disponibles_runtime else []

default_ss_values = {
    'modelos_seleccionados': default_modelos_runtime,
    'tamanio_fuente': "Mediano", # Mantener otros defaults
    'prompt_personalizado': PROMPT_FIJO,
    'limpiar_texto_activo': False,
    'deepseek_temperature': 0.7,
    'deepseek_top_p': 0.9,
    'ultimas_respuestas': {},
    'archivo_procesado_nombre': None
}
for key, value in default_ss_values.items():
    if key not in st.session_state:
        st.session_state[key] = value

# --- Barra lateral para configuraci√≥n ---
with st.sidebar:
    st.header("‚öôÔ∏è Configuraci√≥n")
    archivo = st.file_uploader("üì§ Sube tu archivo", type=["pdf", "docx", "txt"], key="file_uploader")

    # Limpiar resultados si el archivo cambia
    if archivo and st.session_state.get('archivo_procesado_nombre') != archivo.name:
         print(f"Nuevo archivo detectado: {archivo.name}. Limpiando resultados.")
         st.session_state['ultimas_respuestas'] = {}
         st.session_state['archivo_procesado_nombre'] = archivo.name

    st.subheader("ü§ñ Modelos de IA")
    # Filtrar selecci√≥n previa para que solo contenga modelos realmente disponibles ahora
    seleccion_previa = st.session_state.get('modelos_seleccionados', [])
    seleccion_valida = [m for m in seleccion_previa if m in opciones_disponibles_runtime]
    if not seleccion_valida and default_modelos_runtime: # Si la selecci√≥n previa no es v√°lida, usar el default
         seleccion_valida = default_modelos_runtime

    modelos_seleccionados_ui = st.multiselect(
         "Selecciona los modelos a usar:",
         options=opciones_disponibles_runtime, # Usar opciones disponibles
         default=seleccion_valida # Usar default v√°lido
    )
    st.session_state['modelos_seleccionados'] = modelos_seleccionados_ui

    # --- Resto de la Sidebar ---
    st.subheader("üî† Tama√±o de letra")
    tamanio_fuente = st.radio("", ("Peque√±o", "Mediano", "Grande"), index=("Peque√±o", "Mediano", "Grande").index(st.session_state['tamanio_fuente']), key="tamanio_radio", horizontal=True)
    st.session_state['tamanio_fuente'] = tamanio_fuente
    # Aplicar tama√±o de fuente (opcional, el CSS anterior podr√≠a ser suficiente)

    st.subheader("‚öôÔ∏è Opciones de Texto")
    st.session_state['limpiar_texto_activo'] = st.checkbox("Limpiar texto antes del an√°lisis", value=st.session_state['limpiar_texto_activo'])

    if "DeepSeek" in st.session_state['modelos_seleccionados'] and DEEPSEEK_API_KEY:
         st.subheader("üõ†Ô∏è Par√°metros de DeepSeek")
         st.session_state['deepseek_temperature'] = st.slider("Temperatura", 0.0, 1.0, st.session_state['deepseek_temperature'], 0.01, key="ds_temp")
         st.session_state['deepseek_top_p'] = st.slider("Top P", 0.0, 1.0, st.session_state['deepseek_top_p'], 0.01, key="ds_top_p")

    st.subheader("üìù Prompt Personalizado")
    nuevo_prompt = st.text_area("Edita tu Prompt", st.session_state['prompt_personalizado'], height=200, key="prompt_edit")
    col_guardar, col_restaurar = st.columns(2)
    if col_guardar.button("üíæ Guardar Prompt", key="save_prompt"):
        st.session_state['prompt_personalizado'] = nuevo_prompt
        st.success("‚úÖ Prompt guardado.")
    if col_restaurar.button("‚ôªÔ∏è Restaurar Default", key="restore_prompt"):
        st.session_state['prompt_personalizado'] = PROMPT_FIJO
        st.rerun() # Usar st.rerun para refrescar el text_area

    st.markdown("---") # Separador antes del bot√≥n de analizar
    analizar_habilitado = bool(archivo and st.session_state['modelos_seleccionados'])
    analizar = st.button("üöÄ Analizar Archivo", disabled=not analizar_habilitado, type="primary", key="analyze_button")
    if analizar:
         # Guardar el estado de que el bot√≥n fue presionado para el flujo principal
         st.session_state['analizar_boton_presionado_flag'] = True


# --- Contenido Principal ---
col1, col2 = st.columns([1, 2]) # O ajusta como prefieras

with col2:
    st.subheader("üìã Resultados del An√°lisis")
    placeholders = {} # Diccionario para los contenedores de resultados
    respuestas_a_mostrar = {} # Diccionario para los botones de descarga

    # Determinar si ejecutar o mostrar resultados previos
    if st.session_state.get('analizar_boton_presionado_flag'):
        # Ejecutar an√°lisis
        st.session_state['analizar_boton_presionado_flag'] = False # Resetear flag
        if archivo:
            with st.spinner('Analizando documento... ‚è≥'):
                contenido = leer_archivo(archivo)
                if contenido:
                    texto_para_analizar = contenido
                    if st.session_state['limpiar_texto_activo']:
                        texto_para_analizar = limpiar_texto(texto_para_analizar)

                    full_prompt = f"{st.session_state['prompt_personalizado']}\n\n--- Contenido del archivo ---\n{texto_para_analizar}"

                    respuestas_nuevas = {}
                    modelos_a_ejecutar = st.session_state['modelos_seleccionados']

                    # Crear expanders y placeholders ANTES del bucle
                    for modelo_exp in modelos_a_ejecutar:
                         placeholders[modelo_exp] = st.expander(f"Respuesta de {modelo_exp}", expanded=True).empty()

                    # Llamar a las APIs y actualizar placeholders
                    for modelo in modelos_a_ejecutar:
                        respuesta_modelo = None
                        with st.spinner(f"Contactando a {modelo}..."): # Usar st.spinner correctamente
                            try:
                                if modelo == "Gemini":
                                    respuesta_modelo = analizar_con_gemini(full_prompt)
                                elif modelo == "DeepSeek":
                                    respuesta_modelo = analizar_con_deepseek(full_prompt, st.session_state['deepseek_temperature'], st.session_state['deepseek_top_p'])
                                elif modelo.startswith("gpt-") or modelo == "o3-mini" or modelo == "gpt-4o-mini":
                                    respuesta_modelo = analizar_con_openai(full_prompt, model_id=modelo)
                                else:
                                    respuesta_modelo = f"‚ö†Ô∏è Modelo '{modelo}' no implementado."
                            except Exception as e_api:
                                respuesta_modelo = f"‚ö†Ô∏è Error inesperado llamando a {modelo}: {e_api}"

                        # Guardar respuesta
                        respuestas_nuevas[modelo] = respuesta_modelo

                        # Actualizar placeholder correspondiente
                        if isinstance(respuesta_modelo, str) and not respuesta_modelo.startswith("‚ö†Ô∏è"):
                            placeholders[modelo].markdown(respuesta_modelo)
                        elif isinstance(respuesta_modelo, str) and respuesta_modelo.startswith("‚ö†Ô∏è"):
                            placeholders[modelo].error(respuesta_modelo)
                        else:
                            placeholders[modelo].warning(f"Respuesta inesperada o vac√≠a para {modelo}.")

                    # Guardar en estado de sesi√≥n y preparar para mostrar/descargar
                    st.session_state['ultimas_respuestas'] = respuestas_nuevas
                    st.session_state['archivo_procesado_nombre'] = archivo.name
                    respuestas_a_mostrar = respuestas_nuevas

                else:
                    st.error("‚ùó No se pudo leer el archivo correctamente.")
                    st.session_state['ultimas_respuestas'] = {} # Limpiar
                    st.session_state['archivo_procesado_nombre'] = None

        elif not archivo:
            st.error("‚ùó Primero debes subir un archivo para analizarlo.")
            st.session_state['ultimas_respuestas'] = {} # Limpiar
            st.session_state['archivo_procesado_nombre'] = None

    # Mostrar resultados previos si no se analiz√≥ ahora
    elif st.session_state.get('ultimas_respuestas') and st.session_state.get('archivo_procesado_nombre'):
         # No mostrar si el archivo actual no coincide con el procesado
         if archivo and st.session_state.get('archivo_procesado_nombre') == archivo.name:
              st.info("Mostrando resultados del an√°lisis anterior para este archivo.")
              respuestas_a_mostrar = st.session_state['ultimas_respuestas']
              # Recrear expanders para mostrar resultados previos
              for modelo, respuesta in respuestas_a_mostrar.items():
                   with st.expander(f"Respuesta de {modelo} (anterior)", expanded=True):
                        if isinstance(respuesta, str) and not respuesta.startswith("‚ö†Ô∏è"):
                             st.markdown(respuesta)
                        elif isinstance(respuesta, str) and respuesta.startswith("‚ö†Ô∏è"):
                             st.error(respuesta)
                        else:
                             st.warning(f"Respuesta anterior inesperada o vac√≠a para {modelo}.")
         else:
             # Si hay un archivo diferente cargado, no mostrar resultados viejos
             if archivo and st.session_state.get('archivo_procesado_nombre') != archivo.name:
                  st.info("Sube un archivo y presiona 'Analizar Archivo'.")
             # Si no hay archivo cargado, tampoco mostrar
             elif not archivo:
                  st.info("Sube un archivo y presiona 'Analizar Archivo'.")

    # Mensaje inicial si no hay nada que mostrar
    elif not st.session_state.get('ultimas_respuestas'):
         st.info("Sube un archivo y selecciona modelos en la barra lateral, luego presiona 'Analizar Archivo'.")


    # --- Secci√≥n de Descarga (Siempre visible si hay algo que mostrar/descargar) ---
    if respuestas_a_mostrar:
        st.markdown("---")
        st.subheader("Descargar Resultados como DOCX")
        nombre_original_para_descarga = st.session_state.get('archivo_procesado_nombre', "documento")

        # Usar columnas para los botones de descarga si hay muchos
        num_respuestas_validas = sum(1 for resp in respuestas_a_mostrar.values() if isinstance(resp, str) and not resp.startswith("‚ö†Ô∏è"))
        cols_descarga = st.columns(max(1, num_respuestas_validas)) # Crear al menos una columna
        col_idx = 0

        for modelo, respuesta in respuestas_a_mostrar.items():
             if isinstance(respuesta, str) and not respuesta.startswith("‚ö†Ô∏è"):
                 nombre_archivo_docx = generar_nombre_archivo_docx(nombre_original_para_descarga, modelo)
                 if nombre_archivo_docx:
                     bytes_docx = crear_docx_bytes(respuesta) # ¬°Usando placeholder!
                     # Colocar cada bot√≥n en una columna diferente
                     with cols_descarga[col_idx % len(cols_descarga)]:
                          st.download_button(
                              label=f"‚¨áÔ∏è {modelo}", # Etiqueta m√°s corta
                              data=bytes_docx,
                              file_name=nombre_archivo_docx,
                              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                              key=f"download_{modelo}_{nombre_original_para_descarga}"
                          )
                     col_idx += 1
                 # else: # No deber√≠a pasar si el modelo tiene nombre
                 #     with cols_descarga[col_idx % len(cols_descarga)]:
                 #          st.warning(f"No se gener√≥ nombre para {modelo}")
                 #     col_idx += 1


# --- Bot√≥n de finalizar sesi√≥n en Sidebar ---
st.sidebar.markdown("---")
if st.sidebar.button("‚ùå Finalizar Sesi√≥n", key="finish_session"):
    # Definir defaults aqu√≠ para estar seguro del alcance
    default_final_values = {
        'modelos_seleccionados': default_modelos_runtime, 'tamanio_fuente': "Mediano",
        'prompt_personalizado': PROMPT_FIJO, 'limpiar_texto_activo': False,
        'deepseek_temperature': 0.5, 'deepseek_top_p': 0.7,
        'ultimas_respuestas': {}, 'archivo_procesado_nombre': None,
        'analizar_boton_presionado_flag': False # Resetear flag del bot√≥n tambi√©n
    }
    keys_to_reset = list(default_final_values.keys())
    for key in keys_to_reset:
        if key in default_final_values:
              st.session_state[key] = default_final_values[key]
        elif key in st.session_state: # Eliminar otros estados si existen
              del st.session_state[key]
    st.success("‚úÖ Sesi√≥n finalizada.")
    st.rerun() # Usar st.rerun para refrescar todo
